"""File handling utilities for SCIENCEMENTOR."""
import os
from typing import Tuple, Optional

# Maximum file size (5MB)
MAX_FILE_SIZE = 5 * 1024 * 1024

# Maximum content length for LLM context
MAX_CONTENT_LENGTH = 10000


def extract_text_from_file(file_path: str, file_type: str) -> Tuple[bool, str]:
    """Extract text content from a file.
    
    Args:
        file_path: Path to the uploaded file.
        file_type: MIME type or extension of the file.
        
    Returns:
        Tuple of (success, content_or_error_message).
    """
    try:
        # Check file size
        file_size = os.path.getsize(file_path)
        if file_size > MAX_FILE_SIZE:
            return (False, f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB.")
        
        # Handle text files
        if file_type in ['text/plain', '.txt']:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return (True, truncate_content(content))
        
        # Handle PDF files
        if file_type in ['application/pdf', '.pdf']:
            return extract_from_pdf(file_path)
        
        # Handle Word documents
        if file_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', '.docx']:
            return extract_from_docx(file_path)
        
        if file_type in ['application/msword', '.doc']:
            return (False, "Legacy .doc format is not supported. Please convert to .docx")
        
        return (False, f"Unsupported file type: {file_type}")
        
    except Exception as e:
        return (False, f"Error reading file: {str(e)}")


def extract_from_pdf(file_path: str) -> Tuple[bool, str]:
    """Extract text from a PDF file.
    
    Args:
        file_path: Path to the PDF file.
        
    Returns:
        Tuple of (success, content_or_error_message).
    """
    try:
        from PyPDF2 import PdfReader
        
        reader = PdfReader(file_path)
        text_parts = []
        
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
        
        if not text_parts:
            return (False, "Could not extract text from PDF. The file may be scanned or image-based.")
        
        content = "\n\n".join(text_parts)
        return (True, truncate_content(content))
        
    except ImportError:
        return (False, "PDF processing is not available. Please install PyPDF2.")
    except Exception as e:
        return (False, f"Error reading PDF: {str(e)}")


def extract_from_docx(file_path: str) -> Tuple[bool, str]:
    """Extract text from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file.
        
    Returns:
        Tuple of (success, content_or_error_message).
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        if not text_parts:
            return (False, "Could not extract text from document. The file may be empty.")
        
        content = "\n\n".join(text_parts)
        return (True, truncate_content(content))
        
    except ImportError:
        return (False, "DOCX processing is not available. Please install python-docx.")
    except Exception as e:
        return (False, f"Error reading DOCX: {str(e)}")


def truncate_content(content: str) -> str:
    """Truncate content to maximum length for LLM context.
    
    Args:
        content: The text content to truncate.
        
    Returns:
        Truncated content with indicator if truncated.
    """
    if len(content) <= MAX_CONTENT_LENGTH:
        return content
    
    # Truncate and add indicator
    truncated = content[:MAX_CONTENT_LENGTH]
    # Try to break at word boundary
    last_space = truncated.rfind(' ')
    if last_space > MAX_CONTENT_LENGTH - 100:
        truncated = truncated[:last_space]
    
    return truncated + "\n\n[... content truncated due to length ...]"


def get_file_extension(filename: str) -> str:
    """Get file extension from filename.
    
    Args:
        filename: The filename to extract extension from.
        
    Returns:
        Lowercase file extension including the dot.
    """
    if '.' in filename:
        return '.' + filename.rsplit('.', 1)[1].lower()
    return ''
